from dash import Dash, dcc, html, Input, Output, State
#import os
from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_COLOR_INDEX
import copy
import plotly.express as px
import plotly.graph_objects as go
import pandas as pd
import re
#from flask import Flask#, send_from_directory
# from functions.xml_parser import xml_parser
from app.functions.upload_files import file_download_link, save_file, uploaded_files

#server = Flask(__name__)

badwords = {
    'Type 1':[["neat","cool"],'cat1','type1'],
    'Type 2':[["and","or","but"],'cat2', 'type2'],
    'Type 3':[["word","them"],'cat3','type3']
    }

def run_style(runs, c):
    para = {}
    new_dict = {}
    index = 0
    pattern_master = r'\b{}.?.?[.,!:;]?\b'
    for r in runs:
        print("line 29 == run_style: " + r.text,end="\n\n") #+ ".  para is " + str(len(para)) + " runs long")
        styles = {}
        if r.bold:
            styles["font-weight"] = "bold"
        if r.italic:
            styles["font-style"] = "italic"
        if r.underline:
            styles["text-decoration"] = "underline" # https://www.w3schools.com/cssref/pr_text_text-decoration.asp
        if r.font.all_caps:
            styles["font-variant"] = "small-caps"
        if not r.font.highlight_color is None:
            #print(run.font.highlight_color)
            if run.font.highlight_color is WD_COLOR_INDEX.AUTO:
                    #print(run.text + ": AUTO")
                style["background-color"]=""
            if run.font.highlight_color is WD_COLOR_INDEX.BLACK:
                    #print(run.text + ": BLACK")
                style["background-color"]="black"
            if run.font.highlight_color is WD_COLOR_INDEX.BLUE:
                    #print(run.text + ": BLUE")
                style["background-color"]="blue"
            if run.font.highlight_color is WD_COLOR_INDEX.BRIGHT_GREEN:
                    #print(run.text + ": BRIGHT_GREEN")
                style["background-color"]="#00FF00"
            if run.font.highlight_color is WD_COLOR_INDEX.DARK_BLUE:
                    #print(run.text + ": DARK_BLUE")
                style["background-color"]="#000080"
            if run.font.highlight_color is WD_COLOR_INDEX.DARK_RED:
                    #print(run.text + ": DARK_RED")
                style["background-color"]="#800000"
            if run.font.highlight_color is WD_COLOR_INDEX.DARK_YELLOW:
                    #print(run.text + ": DARK_YELLOW")
                style["background-color"]="#808000"
            if run.font.highlight_color is WD_COLOR_INDEX.GRAY_25:
                    #print(run.text + ": GRAY_25")
                style["background-color"]="#C0C0C0"
            if run.font.highlight_color is WD_COLOR_INDEX.GRAY_50:
                    #print(run.text + ": GRAY_50")
                style["background-color"]="#808080"
            if run.font.highlight_color is WD_COLOR_INDEX.GREEN:
                    #print(run.text + ": GREEN")
                style["background-color"]="#008000"
            if run.font.highlight_color is WD_COLOR_INDEX.PINK:
                    #print(run.text + ": PINK")
                style["background-color"]="#FF00FF"
            if run.font.highlight_color is WD_COLOR_INDEX.RED:
                    #print(run.text + ": RED")
                style["background-color"]="#FF0000"
            if run.font.highlight_color is WD_COLOR_INDEX.TEAL:
                    #print(run.text + ": TEAL")
                style["background-color"]="#008080"
            if run.font.highlight_color is WD_COLOR_INDEX.TURQUOISE:
                    #print(run.text + ": TURQUOISE")
                style["background-color"]="#00FFFF"
            if run.font.highlight_color is WD_COLOR_INDEX.VIOLET:
                    #print(run.text + ": VIOLET")
                style["background-color"]="#800080"
            if run.font.highlight_color is WD_COLOR_INDEX.WHITE:
                    #print(run.text + ": WHITE")
                style["background-color"]="#FFFFFF"
            if run.font.highlight_color is WD_COLOR_INDEX.YELLOW:
                    #print(run.text + ": YELLOW")
                style["background-color"]="#FFFF00"
        if not r.font.size is None:
            styles["font-size"] = str(r.font.size.pt) + "pt"
        if r.font.small_caps:
            styles["font-variant"] = "small-caps"

        '''Mark runs as either containing a keyword
        from a particular category or not'''
        for cat in c: # for each category selected
            for word in badwords[cat][0]: # for each keyword
                pattern_str = pattern_master.format(word)
                pattern = re.compile(pattern_str,re.IGNORECASE)
                if pattern.search(r.text): # if the current keyword is in this run
                #if re.search(pattern,r.text,re.IGNORECASE):
                    if r.text in para.keys(): # if the current run has already been eval'd against at least one keyword
                        if cat in para[r.text]["categories"]: # if the current category has already been evaluated
                            #para[r.text]["keywords"][cat].append(word)
                            para[r.text]["hasKeyword"] = True
                        else:
                            para[r.text]["categories"].append(cat)
                            para[r.text]["hasKeyword"] = True
                    else:
                        para[r.text]={
                            "text": r.text,
                            "categories": [cat],
                            "hasKeyword":True,
                            #"words":{w: {}}, #for w in re.split(r' ',r.text.strip())},
                            "style":styles,
                            "className":""
                        }
                else: # if the current keyword isn't in this run
                    if not r.text in para.keys(): # if this run has already been evaluated, do nothing
                        para[r.text]={
                            "text": r.text,
                            "categories": [],
                            "hasKeyword":False,
                            #"words":{w: {} for w in re.split(r' ',r.text.strip())},
                            "style": styles,
                            "className":""
                        }
        print(para[r.text],end=": ")
        print(para[r.text]["categories"],end="\n\n")

    ''' build the flatter dictionary that holds
    all of the runs and split runs to return'''
    return_para={}
    for k, v in para.items(): # keys are runs, values are dictionaries describing runs
        #print(para[k]["hasKeyword"])
        if para[k]["hasKeyword"]: # if there's a keyword to find in the run
            para[k]["words"]={}
            for w in re.split(r' ',para[k]["text"].strip()): # for each word in the run
                for c in para[k]["categories"]: # for each category/[list of keywords] pair:
                    for kw in badwords[c][0]: # for each keyword in that category
                        pattern_str = pattern_master.format(kw) # insert that keyword into the regex
                        pattern = re.compile(pattern_str,re.IGNORECASE)
                        #print("line 153 == reassembly required: " + pattern.pattern)
                        if pattern.match(w): # if the word is a keyword, give it specific style
                            para[k]["words"][w]={
                                "text":[w + " ",html.Div(className="label",children=badwords[c][2])],
                                "style":{},
                                "hasKeyword":True,
                                "keyword":{c:kw},
                                "className": badwords[c][1] + " category"
                            }
                            """print("line 165 == "+w+" | keyword not in current dictionary",end=" ")
                            print(para[k]["words"][w], end = "\n\n")"""
                        else: # if the word w does not match the current keyword kw
                            if w in para[k]["words"].keys(): # if the word has already been accounted for under another keyword
                                """print("line 169 == "+w+" | seent it, not a current keyword",end=" ")
                                print(para[k]["words"][w], end = "\n\n")"""
                                break
                            else: # if the word has not already been noted in the dictionary
                                para[k]["words"][w]={
                                    "text":w + " ",
                                    "style":styles,
                                    "hasKeyword":False,
                                    "kw":{},
                                    "className": ""
                                }
                                """print("line 180 == "+w+" | not a keyword and not in dictionary",end=" ")
                                print(para[k]["words"][w], end = "\n\n")"""
                return_para[k+kw+w]=para[k]["words"][w]
                index+=1
        else:
            return_para[k]=v # append as is
            return_para[k]["text"] += " "

    """print("line 187 == New Dictionary")
    print(return_para)"""
    return [html.Span(children=v["text"],style=v["style"],className=v["className"]) for k,v in return_para.items()]

def para_style(para,cats):
    if para.style.name == "Heading 1":
        return html.H1(children=run_style(para.runs,cats))
    if para.style.name == "Heading 2":
        return html.H2(children=run_style(para.runs,cats))
    if para.style.name == "Normal":
        return html.P(children=run_style(para.runs,cats))
    else:
        print("line 156 == Other " + "".join(cats))
        return html.P(children=run_style(para.runs,cats))

def register_callbacks(app, flask=True):
    print("\n"*5+"~"*30+" New Run "+"~"*30+"\n"*5)
    @app.callback(
        [#Output("switches-input",'options'),
        Output("language_picker",'options')],
        Input("auth-left","children")
    )
    def checklist(child):
        return [list(badwords.keys())]#,list(badwords.keys())

    @app.callback(                          # callback on upload file chosen
        [Output("file-list", "children"), Output("preview-title","children")],
        [Input("upload-data", "filename"), Input("upload-data", "contents")],
    )
    def update_active_files(uploaded_filenames, uploaded_file_contents):
        """Save uploaded files and regenerate the file list."""
        names = []
        if uploaded_filenames is not None and uploaded_file_contents is not None:
            for name, data in zip(uploaded_filenames, uploaded_file_contents):
                save_file(name, data)

        files = uploaded_files()
        if len(files) == 0 or files is None:
            return [html.Li("No files yet!")],'Document Preview'
        else:
            return [html.Li(file_download_link(filename)) for filename in files], [name.split('.')[0] for name in uploaded_filenames]

    @app.callback(
        #[
        Output("text-preview","children"),#Output("bar","figure")],
        [Input("language_picker","value"),Input("ready","n_clicks")]
    )
    def newfunc(checks,n):
        """This function needs to take the language category
        and uploaded file to start returning data"""
        import os
        files = uploaded_files()
        runs=""
        targetwords={}
        """Build the dictionary to hold counts of keywords"""
        for key, v in badwords.items():
            targetwords[key]={}
            for value in v[0]:
                targetwords[key][value]=0

        """Assess the file for keywords and return augmented text for display"""
        for file in files:
            doc=Document(files[file]) # need to figure out how to reference the path for a specific doc
            for paragraph in doc.paragraphs:
                for k, v in badwords.items():
                    for value in v[0]:
                        if value in paragraph.text:
                            targetwords[k][value]+=paragraph.text.count(value)
            os.remove(files[file])         #will need to move once we get the download function up and running
            items = []
            keyword_outline = []
            for k in targetwords:
                keyword_outline.append(html.P(k + ": "))
                keyword_outline.append(html.Ul(
                    children=[html.Li(str(targetwords[k][key1])+ " " + key1 + " keywords.") for key1 in targetwords[k]]
                ))

            for paragraph in doc.paragraphs:
                print("line 209 == " + paragraph.text)
                items.append(para_style(paragraph,list(set(badwords.keys())-set(checks))))

            #df = {key: pd.Series(val) for key, val in targetwords.items() }
            #df2 = pd.DataFrame.from_dict(targetwords, orient="index")
            #print(df)
            #print(df2)

            #df3 = pd.DataFrame([{'Category':{},'Keyword':{},'Value':{}}])
            # for each key in targetwords, concatenate a row
            #print(targetwords)
            return items#, go.Figure(px.bar(df))
